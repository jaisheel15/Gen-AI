import os
import docx

def create_dir_if_not_exists(path):
    os.makedirs(path, exist_ok=True)

def generate_standard_contract(filepath):
    doc = docx.Document()
    
    doc.add_heading('MUTUAL BUSINESS PARTNERSHIP AGREEMENT', 0)
    
    doc.add_paragraph('This Mutual Business Partnership Agreement (the "Agreement") is entered into as of June 11, 2026, by and between Alpha Corp, a Delaware corporation ("Alpha"), and Beta Holdings LLC, a Delaware limited liability company ("Beta").')
    
    doc.add_heading('Section 1. Indemnification', level=1)
    doc.add_paragraph('Each party (the "Indemnifying Party") shall indemnify, defend, and hold harmless the other party and its officers, directors, and employees (the "Indemnified Party") from and against any and all third-party claims, losses, damages, liabilities, costs, and expenses (including reasonable attorneys\' fees) arising out of or relating to a material breach of this Agreement, gross negligence, or willful misconduct by the Indemnifying Party in connection with its obligations under this Agreement.')
    
    doc.add_heading('Section 2. Limitation of Liability', level=1)
    doc.add_paragraph('Except for breaches of confidentiality under Section 7, indemnification obligations under Section 1, or willful misconduct, neither party shall be liable for any indirect, incidental, special, consequential, or punitive damages, or loss of profits or revenues under this Agreement. Each party\'s maximum aggregate liability under this Agreement shall be limited to the total fees paid or payable under this Agreement in the twelve (12) months preceding the event giving rise to liability.')
    
    doc.add_heading('Section 3. Governing Law and Jurisdiction', level=1)
    doc.add_paragraph('This Agreement shall be governed by, and construed in accordance with, the laws of the State of Delaware, without regard to its conflict of laws principles. The parties submit to the exclusive jurisdiction of the state and federal courts located in Wilmington, Delaware, for the resolution of any dispute arising hereunder.')
    
    doc.add_heading('Section 4. Term and Termination', level=1)
    doc.add_paragraph('This Agreement shall commence on the Effective Date and continue for one (1) year. Either party may terminate this Agreement for convenience upon ninety (90) days\' written notice to the other party. Either party may terminate this Agreement for cause if the other party materially breaches this Agreement and fails to cure such breach within thirty (30) days of receiving written notice thereof.')
    
    doc.add_heading('Section 5. Intellectual Property Rights', level=1)
    doc.add_paragraph('Each party retains all right, title, and interest in its pre-existing intellectual property. Any intellectual property created or developed solely or jointly by a party in connection with the services shall be owned by the party that created it, and to the extent it constitutes deliverables, shall be licensed to the other party on a non-exclusive, worldwide, royalty-free, perpetual basis.')
    
    doc.add_heading('Section 6. Payment Terms and Invoicing', level=1)
    doc.add_paragraph('Client shall pay all undisputed invoices within thirty (30) days of receipt. Any late payments not subject to a good faith dispute shall accrue interest at a rate of 1.0% per month, or the maximum rate permitted by law, whichever is lower.')
    
    doc.add_heading('Section 7. Confidentiality Obligations', level=1)
    doc.add_paragraph('Each party agrees to hold the other party\'s Confidential Information in strict confidence and to use at least the same degree of care it uses to protect its own confidential information, but in no event less than a reasonable degree of care. Confidential Information shall not be disclosed to third parties except to employees, consultants, or advisors with a need to know. This obligation survives termination for a period of three (3) years.')
    
    doc.save(filepath)
    print(f"Generated standard contract at: {filepath}")

def generate_poison_pill_contract(filepath):
    doc = docx.Document()
    
    doc.add_heading('UNILATERAL VENDOR SERVICES AGREEMENT', 0)
    
    doc.add_paragraph('This Services Agreement (the "Agreement") is entered into as of June 11, 2026, by and between Titan Vendor Inc. ("Vendor") and Client Corp ("Client").')
    
    doc.add_heading('Section 1. Indemnification Obligations', level=1)
    doc.add_paragraph('Client agrees to indemnify, defend, and hold harmless Vendor from any and all third-party claims, damages, losses, or liabilities, including those arising from Vendor\'s own gross negligence, active fault, or willful misconduct under this Agreement. Vendor offers no indemnification of any kind to Client, and Client waives all claims related to third-party actions against Vendor.')
    
    doc.add_heading('Section 2. Limitation of Liability and Damages Cap', level=1)
    doc.add_paragraph('In no event shall Vendor\'s total, aggregate liability under this Agreement exceed Ten Dollars ($10.00) under any theory of law. Client\'s liability to Vendor shall be uncapped and unlimited, and Client shall be liable for all indirect, consequential, punitive, and special damages, including loss of potential business, arising out of any breach.')
    
    doc.add_heading('Section 3. Governing Law and Choice of Forum', level=1)
    doc.add_paragraph('This Agreement shall be governed exclusively by the laws of Pyongyang, North Korea. The parties submit to the exclusive jurisdiction of the military tribunals of Pyongyang, North Korea, for all disputes, and Client waives any right to claim forum non conveniens.')
    
    doc.add_heading('Section 4. Termination and Cure Periods', level=1)
    doc.add_paragraph('Vendor may terminate this Agreement immediately or upon twenty-four (24) hours\' written notice to Client for any reason or no reason. Client has no right to terminate this Agreement for convenience, and may only terminate for material breach if Vendor fails to cure such breach within three hundred and sixty-five (365) days of receipt of written notice thereof.')
    
    doc.add_heading('Section 5. Assignment of Intellectual Property', level=1)
    doc.add_paragraph('Client hereby irrevocably assigns, transfers, and conveys to Vendor all right, title, and interest in and to all pre-existing intellectual property, proprietary software, data, patents, and trademarks owned by Client prior to this Agreement. Any new intellectual property developed under this Agreement shall belong solely to Vendor.')
    
    doc.add_heading('Section 6. Payment Schedule and Interest Penalty', level=1)
    doc.add_paragraph('Client shall pay all invoices within twenty-four (24) hours of receipt. Any late payments shall accrue interest at a rate of 15% per day, compounded hourly, and Vendor shall suspend all services without liability.')
    
    doc.add_heading('Section 7. Confidentiality and Data Exposure', level=1)
    doc.add_paragraph('Client shall keep all information related to Vendor strictly confidential in perpetuity. Vendor has no confidentiality obligations and may disclose Client\'s proprietary data, business plans, and trade secrets to any third party for marketing or other commercial purposes without notice or compensation.')
    
    doc.save(filepath)
    print(f"Generated poison pill contract at: {filepath}")

def generate_consulting_contract(filepath):
    doc = docx.Document()
    
    doc.add_heading('CONSULTING SERVICES AGREEMENT', 0)
    
    doc.add_paragraph('This Consulting Services Agreement (the "Agreement") is entered into as of June 11, 2026, by and between Apex Advisory ("Consultant") and Nexus Enterprises ("Client").')
    
    doc.add_heading('Section 1. Indemnity', level=1)
    doc.add_paragraph('Consultant shall indemnify Client from claims arising out of Consultant\'s gross negligence or willful misconduct in the performance of services. Client shall indemnify Consultant from claims arising out of Client\'s breach of Agreement.')
    
    doc.add_heading('Section 2. Limitation of Liability', level=1)
    doc.add_paragraph('Except for breach of confidentiality, each party\'s total liability under this Agreement shall be limited to the total fees paid by Client to Consultant under this Agreement.')
    
    doc.add_heading('Section 3. Governing Law', level=1)
    doc.add_paragraph('This Agreement shall be governed by, and construed in accordance with, the laws of the State of California. Any disputes shall be settled in the courts of San Francisco, California.')
    
    doc.add_heading('Section 4. Termination', level=1)
    doc.add_paragraph('Either party may terminate this Agreement upon thirty (30) days\' written notice to the other party. Upon termination, Client shall pay Consultant for all services rendered up to the date of termination.')
    
    doc.add_heading('Section 5. IP Ownership', level=1)
    doc.add_paragraph('Upon full payment of all fees, all deliverables created for Client shall become the sole property of Client. Consultant retains ownership of all background consulting tools and methodologies used in performing the services.')
    
    doc.add_heading('Section 6. Payments and Rates', level=1)
    doc.add_paragraph('Client shall pay all invoices within thirty (30) days of receipt. Interest on late payments shall accrue at 0.5% per month.')
    
    doc.add_heading('Section 7. Confidentiality', level=1)
    doc.add_paragraph('Both parties agree to protect the other\'s confidential business information. This obligation shall survive the termination of this agreement for a period of two (2) years.')
    
    doc.save(filepath)
    print(f"Generated consulting contract at: {filepath}")

if __name__ == "__main__":
    dir_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "sample_contracts")
    create_dir_if_not_exists(dir_path)
    
    generate_standard_contract(os.path.join(dir_path, "standard_partnership_agreement.docx"))
    generate_poison_pill_contract(os.path.join(dir_path, "poison_pill_vendor_contract.docx"))
    generate_consulting_contract(os.path.join(dir_path, "consulting_service_agreement.docx"))
